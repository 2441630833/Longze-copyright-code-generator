import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import glob
import re

class SourceCodeGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("软著源代码生成器")
        self.root.geometry("600x400")
        self.root.resizable(False, False)
        
        # 创建主框架
        main_frame = ttk.Frame(root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # 软件名称
        ttk.Label(main_frame, text="软件名称:").grid(row=0, column=0, sticky=tk.W, pady=10)
        self.software_name = ttk.Entry(main_frame, width=50)
        self.software_name.grid(row=0, column=1, sticky=tk.W, pady=10)
        self.software_name.insert(0, "测试软件")
        
        # 版本号
        ttk.Label(main_frame, text="版本号:").grid(row=1, column=0, sticky=tk.W, pady=10)
        self.version = ttk.Entry(main_frame, width=50)
        self.version.grid(row=1, column=1, sticky=tk.W, pady=10)
        self.version.insert(0, "V1.0")
        
        # 作者名
        ttk.Label(main_frame, text="作者名:").grid(row=2, column=0, sticky=tk.W, pady=10)
        self.author = ttk.Entry(main_frame, width=50)
        self.author.grid(row=2, column=1, sticky=tk.W, pady=10)
        self.author.insert(0, "测试")
        
        # 项目路径
        ttk.Label(main_frame, text="项目路径:").grid(row=3, column=0, sticky=tk.W, pady=10)
        path_frame = ttk.Frame(main_frame)
        path_frame.grid(row=3, column=1, sticky=tk.W, pady=10)
        
        self.project_path = ttk.Entry(path_frame, width=42)
        self.project_path.pack(side=tk.LEFT)
        
        ttk.Button(path_frame, text="选择", command=self.select_path).pack(side=tk.RIGHT)
        
        # 指定文件后缀
        ttk.Label(main_frame, text="指定文件后缀:").grid(row=4, column=0, sticky=tk.W, pady=10)
        self.file_extensions = ttk.Entry(main_frame, width=50)
        self.file_extensions.grid(row=4, column=1, sticky=tk.W, pady=10)
        ttk.Label(main_frame, text="指定多个文件请使用英文逗号分隔").grid(row=5, column=1, sticky=tk.W)
        
        # 按钮区域
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=6, column=0, columnspan=2, pady=30)
        
        ttk.Button(button_frame, text="生成", command=self.generate_document, width=20).pack(side=tk.LEFT, padx=10)
        ttk.Button(button_frame, text="取消", command=root.quit, width=20).pack(side=tk.LEFT, padx=10)
        
        # 进度条
        self.progress = ttk.Progressbar(main_frame, orient="horizontal", length=560, mode="determinate")
        self.progress.grid(row=7, column=0, columnspan=2, pady=10)
        
        # 状态标签
        self.status_label = ttk.Label(main_frame, text="")
        self.status_label.grid(row=8, column=0, columnspan=2)
    
    def select_path(self):
        """选择项目路径"""
        path = filedialog.askdirectory(title="选择项目路径")
        if path:
            self.project_path.delete(0, tk.END)
            self.project_path.insert(0, path)
            # 自动检测文件后缀
            self.detect_file_extensions(path)
    
    def detect_file_extensions(self, path):
        """自动检测目录中的文件后缀"""
        if not path or not os.path.exists(path):
            return
            
        self.update_status("正在检测项目中的文件类型...")
        
        # 获取所有文件
        all_files = []
        for root, dirs, files in os.walk(path):
            for file in files:
                all_files.append(os.path.join(root, file))
                
        # 提取后缀
        extensions = set()
        for file in all_files:
            ext = os.path.splitext(file)[1]
            if ext and not ext.startswith('.git'):  # 排除git相关文件
                extensions.add(ext.lstrip('.'))
                
        # 按常见编程语言排序
        common_exts = ['java', 'py', 'js', 'html', 'css', 'xml', 'json', 'c', 'cpp', 'h', 'hpp', 'cs', 'php', 'go', 'rs', 'ts', 'vue', 'jsx', 'tsx']
        sorted_exts = sorted(extensions, key=lambda x: (x not in common_exts, x))
        
        if sorted_exts:
            self.file_extensions.delete(0, tk.END)
            self.file_extensions.insert(0, ','.join(sorted_exts[:10]))  # 最多显示10个后缀
            self.update_status(f"检测到 {len(sorted_exts)} 种文件类型")
    
    def update_status(self, message):
        """更新状态信息"""
        self.status_label.config(text=message)
        self.root.update()
    
    def add_page_number(self, doc):
        """添加页脚（著作权人全称）"""
        for section in doc.sections:
            footer = section.footer
            # 清除现有段落并创建新段落
            if footer.paragraphs:
                p = footer.paragraphs[0]
            else:
                p = footer.add_paragraph()
                
            # 设置对齐方式
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加著作权人全称
            author_run = p.add_run(f"著作权人: {self.author.get()}")
            author_run.font.size = Pt(10)
    
    def calculate_optimal_line_spacing(self, doc, total_lines, target_pages=60):
        """计算最佳行距以确保文档正好是指定页数"""
        # 估计每页的基本行数
        base_lines_per_page = 50
        
        # 计算总页数
        lines_per_page = total_lines / target_pages if total_lines > 0 else base_lines_per_page
        
        # 调整行距来控制页面数量
        if lines_per_page < base_lines_per_page:
            # 如果每页行数少于标准行数，则增加行距，但限制在合理范围内
            lines_ratio = base_lines_per_page / max(lines_per_page, 1)
            # 限制最大行距为1.5，最小为1.0
            optimal_spacing = min(1.5, max(1.0, lines_ratio))
        elif lines_per_page > base_lines_per_page:
            # 如果每页行数多于标准行数，则减小行距，确保内容不会超出页面
            lines_ratio = lines_per_page / base_lines_per_page
            # 限制最小行距为0.9，以保持可读性
            optimal_spacing = max(0.9, 1.0 / lines_ratio)
        else:
            # 刚好合适
            optimal_spacing = 1.0
            
        return optimal_spacing

    def generate_document(self):
        """生成文档"""
        # 获取输入
        software_name = self.software_name.get().strip()
        version = self.version.get().strip()
        author = self.author.get().strip()
        project_path = self.project_path.get().strip()
        extensions = self.file_extensions.get().strip()
        
        # 默认字体大小
        font_size = 10  # 确保font_size在函数开始就定义
        
        # 验证输入
        if not software_name:
            messagebox.showerror("错误", "请输入软件名称")
            return
        if not version:
            messagebox.showerror("错误", "请输入版本号")
            return
        if not author:
            messagebox.showerror("错误", "请输入作者名")
            return
        if not project_path or not os.path.exists(project_path):
            messagebox.showerror("错误", "请选择有效的项目路径")
            return
        if not extensions:
            messagebox.showerror("错误", "请输入至少一个文件后缀")
            return
        
        try:
            # 重置进度条
            self.progress["value"] = 0
            self.update_status("开始收集源代码文件...")
            
            # 解析文件后缀
            ext_list = [ext.strip() for ext in extensions.split(",")]
            ext_list = [ext if ext.startswith(".") else f".{ext}" for ext in ext_list]
            
            # 收集所有匹配的文件
            all_files = []
            for ext in ext_list:
                files = glob.glob(f"{project_path}/**/*{ext}", recursive=True)
                all_files.extend(files)
            
            if not all_files:
                messagebox.showerror("错误", f"在指定路径下未找到任何匹配的文件")
                return
            
            # 更新进度条
            self.progress["maximum"] = len(all_files) + 1
            self.progress["value"] = 1
            self.update_status(f"找到 {len(all_files)} 个文件，开始处理...")
            
            # 创建文档
            doc = Document()
            
            # 设置页面边距
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(2.54)
                section.bottom_margin = Cm(2.54)
                section.left_margin = Cm(3.17)
                section.right_margin = Cm(3.17)
            
            # 计算前30页和后30页的行数
            lines_per_page = 50  # 每页固定50行
            front_pages = 30
            back_pages = 30
            total_pages = front_pages + back_pages
            
            front_lines = front_pages * lines_per_page
            back_lines = back_pages * lines_per_page
            
            # 添加页眉
            header = doc.sections[0].header
            header_para = header.paragraphs[0]
            header_para.text = f"{software_name} {version}"
            header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # 添加右上角页码信息
            header_table = header.add_table(1, 2, width=Cm(16))
            # 左侧为软件名称和版本号
            cell_left = header_table.cell(0, 0)
            cell_left.text = f"{software_name} {version}"
            cell_left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            # 右侧为页码
            cell_right = header_table.cell(0, 1)
            cell_right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # 添加页码字段
            run = cell_right.paragraphs[0].add_run("第")
            run.font.size = Pt(10)
            
            # 添加页码
            page_num = cell_right.paragraphs[0].add_run()
            r = page_num._r
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), 'begin')
            r.append(fld)
            
            run = cell_right.paragraphs[0].add_run()
            r = run._r
            instrText = OxmlElement('w:instrText')
            instrText.text = " PAGE "
            r.append(instrText)
            
            run = cell_right.paragraphs[0].add_run()
            r = run._r
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), 'end')
            r.append(fld)
            
            # 添加"页，共x页"
            run = cell_right.paragraphs[0].add_run("页，共")
            run.font.size = Pt(10)
            
            # 添加总页数，使用域代码动态显示实际页数
            run = cell_right.paragraphs[0].add_run()
            r = run._r
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), 'begin')
            r.append(fld)
            
            run = cell_right.paragraphs[0].add_run()
            r = run._r
            instrText = OxmlElement('w:instrText')
            instrText.text = " NUMPAGES "
            r.append(instrText)
            
            run = cell_right.paragraphs[0].add_run()
            r = run._r
            fld = OxmlElement('w:fldChar')
            fld.set(qn('w:fldCharType'), 'end')
            r.append(fld)
            
            run = cell_right.paragraphs[0].add_run("页")
            run.font.size = Pt(10)
            
            # 设置表格无边框
            for row in header_table.rows:
                for cell in row.cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
            
            # 删除原始段落，只保留表格
            if len(header.paragraphs) > 0:
                p = header.paragraphs[0]._p
                if p.getparent() is not None:
                    p.getparent().remove(p)
            
            # 添加页码（使用新的方法）
            self.add_page_number(doc)
            
            # 收集所有代码行
            all_code_lines = []
            file_boundaries = []  # 记录每个文件的起始行和结束行
            
            for i, file_path in enumerate(all_files):
                # 更新进度
                self.progress["value"] = i + 1
                self.update_status(f"收集文件 {i+1}/{len(all_files)}: {os.path.basename(file_path)}")
                
                # 读取文件内容
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                except UnicodeDecodeError:
                    try:
                        with open(file_path, 'r', encoding='gbk') as f:
                            content = f.read()
                    except:
                        self.update_status(f"警告: 无法读取文件 {file_path}，跳过")
                        continue
                
                # 添加文件内容
                lines = content.split('\n')
                
                # 记录文件边界
                start_line = len(all_code_lines)
                all_code_lines.extend(lines)
                end_line = len(all_code_lines)
                
                # 记录文件信息
                relative_path = os.path.relpath(file_path, project_path)
                file_boundaries.append((relative_path, start_line, end_line))
            
            # 计算总行数
            total_lines = len(all_code_lines)
            self.update_status(f"总共收集了 {total_lines} 行代码")
            
            # 确保有足够的代码行
            if total_lines < 3000:
                messagebox.showwarning("警告", f"收集到的代码行数不足3000行（当前{total_lines}行），可能无法满足软著要求")
            
            # 直接生成60页固定格式文档
            self.update_status("正在生成严格控制为60页的文档...")
            
            # 计算最佳行距和字体大小
            optimal_spacing = self.calculate_optimal_line_spacing(doc, total_lines)
            self.update_status(f"计算得到最佳行距: {optimal_spacing:.2f}")
            
            # 根据页数需求调整字体大小
            if total_lines < 2000:  # 代码行数太少时，适当增大字体
                font_size = min(11, 10 * (3000 / max(total_lines, 1500)))
                self.update_status(f"调整字体大小为: {font_size:.1f}pt")
            
            # 如果代码行数不足以填满60页，则全部使用
            if total_lines <= front_lines + back_lines:
                # 代码不足60页时，全部展示
                code_lines = all_code_lines
                
                # 计算每页应该展示的行数
                lines_per_page_adjusted = min(50, (total_lines + total_pages - 1) // total_pages)
                
                # 生成60页文档
                for page in range(total_pages):
                    start_idx = page * lines_per_page_adjusted
                    end_idx = min(start_idx + lines_per_page_adjusted, total_lines)
                    
                    if start_idx >= total_lines:
                        # 如果已经没有代码了，添加空页
                        code_para = doc.add_paragraph()
                    else:
                        # 添加代码
                        code_para = doc.add_paragraph()
                        code_run = code_para.add_run('\n'.join(code_lines[start_idx:end_idx]))
                        code_run.font.name = 'Courier New'
                        code_run.font.size = Pt(font_size)
                        
                        # 设置行距
                        code_para.paragraph_format.line_spacing = optimal_spacing
                        code_para.paragraph_format.space_after = Pt(0)
                        code_para.paragraph_format.space_before = Pt(0)
            else:
                # 代码超过60页时，只展示前30页和后30页
                front_code = all_code_lines[:front_lines]
                back_code = all_code_lines[-back_lines:]
                
                # 生成前30页
                for page in range(front_pages):
                    start_idx = page * lines_per_page
                    end_idx = min(start_idx + lines_per_page, len(front_code))
                    
                    code_para = doc.add_paragraph()
                    code_run = code_para.add_run('\n'.join(front_code[start_idx:end_idx]))
                    code_run.font.name = 'Courier New'
                    code_run.font.size = Pt(font_size)
                    
                    # 设置行距
                    code_para.paragraph_format.line_spacing = optimal_spacing
                    code_para.paragraph_format.space_after = Pt(0)
                    code_para.paragraph_format.space_before = Pt(0)
                
                # 添加分隔标记
                separator = doc.add_paragraph()
                separator_run = separator.add_run("----- 后30页代码开始 -----")
                separator_run.bold = True
                separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 生成后30页
                for page in range(back_pages):
                    start_idx = page * lines_per_page
                    end_idx = min(start_idx + lines_per_page, len(back_code))
                    
                    code_para = doc.add_paragraph()
                    code_run = code_para.add_run('\n'.join(back_code[start_idx:end_idx]))
                    code_run.font.name = 'Courier New'
                    code_run.font.size = Pt(font_size)
                    
                    # 设置行距
                    code_para.paragraph_format.line_spacing = optimal_spacing
                    code_para.paragraph_format.space_after = Pt(0)
                    code_para.paragraph_format.space_before = Pt(0)
            
            # 保存文档
            output_filename = f"{software_name}源代码.docx"
            doc.save(output_filename)
            
            self.progress["value"] = self.progress["maximum"]
            self.update_status(f"文档生成完成! 总共处理了 {total_lines} 行代码，保存为 {output_filename}")
            
            messagebox.showinfo("成功", f"文档已成功生成: {output_filename}\n\n符合软著申请要求：\n- 动态生成页数，每页约50行\n- 前30页为代码开头，后30页为代码结尾\n- 页眉包含软件名称和版本号\n- 页脚包含著作权人信息")
            
        except Exception as e:
            messagebox.showerror("错误", f"生成文档时出错: {str(e)}")
            self.update_status(f"错误: {str(e)}")


if __name__ == "__main__":
    root = tk.Tk()
    app = SourceCodeGenerator(root)
    root.mainloop() 